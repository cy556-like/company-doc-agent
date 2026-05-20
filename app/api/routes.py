"""
FastAPI 路由定义
支持流式输出（SSE）
"""
import os
import shutil
import json
import asyncio
from typing import Optional

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from app.agent.core import chat, chat_stream_generator, TOOL_DISPLAY_NAMES
from app.rag.document import index_document, search_documents, list_indexed_documents, read_document_content
from app.memory.manager import (
    get_history_messages, clear_session_history,
    create_chat, list_chats, delete_chat, rename_chat, update_chat_time,
)
from app.auth.user_manager import login_user, register_user
from app.config import settings, AVAILABLE_MODELS, set_current_model, get_current_model

router = APIRouter()


# ===== 请求/响应模型 =====
class ChatRequest(BaseModel):
    """聊天请求"""
    message: str
    session_id: str = "default"


class ChatResponse(BaseModel):
    """聊天响应"""
    response: str
    session_id: str


class SearchRequest(BaseModel):
    """文档搜索请求"""
    query: str
    top_k: int = 3


class AuthRequest(BaseModel):
    """认证请求"""
    username: str
    password: str


class RenameChatRequest(BaseModel):
    """重命名会话请求"""
    username: str
    chat_id: str
    new_title: str


class SetModelRequest(BaseModel):
    """切换模型请求"""
    model_id: str


# ===== 认证接口 =====

@router.post("/auth/login", summary="用户登录")
async def auth_login(req: AuthRequest):
    """用户登录验证"""
    result = login_user(req.username, req.password)
    return result


@router.post("/auth/register", summary="用户注册")
async def auth_register(req: AuthRequest):
    """用户注册"""
    result = register_user(req.username, req.password)
    return result


# ===== 会话管理接口 =====

@router.post("/chats", summary="创建新会话")
async def create_new_chat(username: str = Query(..., description="用户名"), title: str = Query("新对话", description="会话标题")):
    """为用户创建一个新的聊天会话"""
    chat_info = create_chat(username, title)
    return {"success": True, "chat": chat_info}


@router.get("/chats", summary="获取用户的所有会话")
async def get_user_chats(username: str = Query(..., description="用户名")):
    """获取用户的所有聊天会话列表"""
    chats = list_chats(username)
    return {"success": True, "chats": chats}


@router.delete("/chats/{chat_id}", summary="删除会话")
async def delete_user_chat(chat_id: str, username: str = Query(..., description="用户名")):
    """删除指定的聊天会话"""
    delete_chat(username, chat_id)
    return {"success": True, "message": "会话已删除"}


@router.put("/chats/{chat_id}/rename", summary="重命名会话")
async def rename_user_chat(chat_id: str, req: RenameChatRequest):
    """重命名指定的聊天会话"""
    rename_chat(req.username, req.chat_id, req.new_title)
    return {"success": True, "message": "会话已重命名"}


# ===== 核心 Agent 接口 =====

@router.post("/chat", response_model=ChatResponse, summary="与 Agent 对话（非流式）")
async def chat_api(req: ChatRequest):
    """
    非流式对话接口（保留兼容）
    """
    try:
        response = chat(req.message, req.session_id)
        # 更新会话时间
        parts = req.session_id.rsplit("_", 1)
        if len(parts) == 2:
            try:
                update_chat_time(parts[0], req.session_id)
            except Exception:
                pass
        return ChatResponse(response=response, session_id=req.session_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Agent 处理失败: {str(e)}")


@router.post("/chat/stream", summary="与 Agent 对话（流式 SSE）")
async def chat_stream_api(req: ChatRequest):
    """
    流式对话：逐token输出，前端实时显示
    返回 Server-Sent Events (SSE) 格式
    """
    async def generate():
        async for chunk in chat_stream_generator(req.message, req.session_id):
            yield f"data: {json.dumps(chunk, ensure_ascii=False)}\n\n"
        # 更新会话时间
        parts = req.session_id.rsplit("_", 1)
        if len(parts) == 2:
            try:
                update_chat_time(parts[0], req.session_id)
            except Exception:
                pass

    return StreamingResponse(generate(), media_type="text/event-stream")


@router.post("/chat-with-file", summary="ChatGPT风格：文件+消息对话（非流式）")
async def chat_with_file(
    file: UploadFile = File(...),
    message: str = Form(...),
    session_id: str = Form("default"),
):
    """
    非流式文件对话接口（保留兼容）
    """
    # 检查文件格式
    allowed_ext = {".pdf", ".txt", ".docx"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")

    # 保存文件到知识库目录
    file_path = os.path.join(settings.DOCUMENTS_DIR, file.filename)
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        # 尝试索引到向量库（失败不影响对话）
        try:
            index_document(file_path, file.filename)
        except Exception:
            pass

        # 提取文件内容
        try:
            file_content = read_document_content(file_path)
        except Exception as e:
            file_content = f"(无法读取文件内容: {str(e)})"

        # 限制文件内容长度
        max_chars = 8000
        if len(file_content) > max_chars:
            file_content = file_content[:max_chars] + f"\n...(已截断，共{len(file_content)}字符)"

        # 构建增强消息
        enhanced_message = f"""[用户上传了文件: {file.filename}]

文件内容如下：
---
{file_content}
---

文件保存路径: {file_path}

用户的问题/要求: {message}"""

        response = chat(enhanced_message, session_id)

        # 更新会话时间
        parts = session_id.rsplit("_", 1)
        if len(parts) == 2:
            try:
                update_chat_time(parts[0], session_id)
            except Exception:
                pass

        return {"response": response, "session_id": session_id}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"处理失败: {str(e)}")


@router.post("/chat-with-file/stream", summary="ChatGPT风格：文件+消息对话（流式 SSE）")
async def chat_with_file_stream(
    file: UploadFile = File(...),
    message: str = Form(...),
    session_id: str = Form("default"),
):
    """
    流式文件对话：文件+消息，逐token输出
    """
    # 检查文件格式
    allowed_ext = {".pdf", ".txt", ".docx"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")

    # 保存文件
    file_path = os.path.join(settings.DOCUMENTS_DIR, file.filename)
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        # 尝试索引
        try:
            index_document(file_path, file.filename)
        except Exception:
            pass

        # 提取文件内容
        try:
            file_content = read_document_content(file_path)
        except Exception as e:
            file_content = f"(无法读取文件内容: {str(e)})"

        max_chars = 8000
        if len(file_content) > max_chars:
            file_content = file_content[:max_chars] + f"\n...(已截断，共{len(file_content)}字符)"

        enhanced_message = f"""[用户上传了文件: {file.filename}]

文件内容如下：
---
{file_content}
---

文件保存路径: {file_path}

用户的问题/要求: {message}"""

        async def generate():
            async for chunk in chat_stream_generator(enhanced_message, session_id):
                yield f"data: {json.dumps(chunk, ensure_ascii=False)}\n\n"
            # 更新会话时间
            parts = session_id.rsplit("_", 1)
            if len(parts) == 2:
                try:
                    update_chat_time(parts[0], session_id)
                except Exception:
                    pass

        return StreamingResponse(generate(), media_type="text/event-stream")

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"处理失败: {str(e)}")


# ===== 文档管理接口 =====

@router.post("/upload", summary="上传文档到知识库")
async def upload_document(file: UploadFile = File(...)):
    """
    上传文档并自动索引到向量数据库
    支持 PDF、TXT、DOCX 格式
    """
    allowed_ext = {".pdf", ".txt", ".docx"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}，仅支持 {allowed_ext}")

    file_path = os.path.join(settings.DOCUMENTS_DIR, file.filename)
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        result = index_document(file_path, file.filename)
        return {"status": "success", "detail": result}
    except Exception as e:
        os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"文档索引失败: {str(e)}")


@router.post("/modify-document", summary="修改文档")
async def modify_document(
    file: UploadFile = File(...),
    instruction: str = Form(...),
    username: str = Form("default"),
):
    """上传文档并根据修改要求进行修改，返回修改后的文件下载链接"""
    allowed_ext = {".pdf", ".txt", ".docx"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}，仅支持 {allowed_ext}")

    temp_dir = os.path.join(settings.DATA_DIR, "temp")
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, file.filename)
    with open(temp_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        content = read_document_content(temp_path)

        from langchain_openai import ChatOpenAI
        from langchain_core.messages import SystemMessage, HumanMessage

        llm = ChatOpenAI(
            api_key=settings.LLM_API_KEY,
            base_url=settings.LLM_BASE_URL,
            model=settings.LLM_MODEL,
            temperature=0.3,
        )

        system_prompt = "你是一个文档修改助手。按修改要求修改文档，只返回修改后的完整内容，不要解释。保持原文档格式和结构，用中文输出。"
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=f"原始文档内容：\n\n{content}\n\n修改要求：{instruction}"),
        ]

        response = llm.invoke(messages)
        modified_content = response.content

        static_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static")
        modified_dir = os.path.join(static_dir, "modified")
        os.makedirs(modified_dir, exist_ok=True)

        output_filename = f"modified_{file.filename}"
        output_path = os.path.join(modified_dir, output_filename)

        if ext == ".docx":
            try:
                from docx import Document
                doc = Document(temp_path)
                for paragraph in doc.paragraphs:
                    paragraph.text = ""
                paragraphs = modified_content.split("\n")
                if doc.paragraphs:
                    doc.paragraphs[0].text = paragraphs[0] if paragraphs else ""
                for p_text in paragraphs[1:]:
                    doc.add_paragraph(p_text)
                doc.save(output_path)
            except ImportError:
                output_filename = f"modified_{os.path.splitext(file.filename)[0]}.txt"
                output_path = os.path.join(modified_dir, output_filename)
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(modified_content)
        elif ext == ".pdf":
            from app.utils.pdf_generator import generate_pdf
            success, actual_path = generate_pdf(modified_content, output_path)
            output_filename = os.path.basename(actual_path)
        else:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(modified_content)

        os.remove(temp_path)
        download_url = f"/static/modified/{output_filename}"

        return {
            "success": True,
            "message": "文档修改完成！",
            "download_url": download_url,
            "filename": output_filename,
        }

    except Exception as e:
        if os.path.exists(temp_path):
            os.remove(temp_path)
        raise HTTPException(status_code=500, detail=f"文档修改失败: {str(e)}")


@router.post("/search", summary="搜索文档内容")
async def search_api(req: SearchRequest):
    """在文档库中搜索相关内容"""
    results = search_documents(req.query, req.top_k)
    return {"query": req.query, "results": results}


@router.get("/documents", summary="列出所有已索引文档")
async def list_documents():
    """获取知识库中所有文档列表"""
    docs = list_indexed_documents()
    return {"documents": docs, "count": len(docs)}


# ===== 对话历史接口 =====

@router.get("/history/{session_id}", summary="获取对话历史")
async def get_history(session_id: str):
    """获取指定会话的对话历史"""
    messages = get_history_messages(session_id)
    return {"session_id": session_id, "messages": messages, "count": len(messages)}


@router.delete("/history/{session_id}", summary="清除对话历史")
async def delete_history(session_id: str):
    """清除指定会话的对话历史"""
    clear_session_history(session_id)
    return {"status": "success", "message": f"会话 {session_id} 的历史已清除"}


# ===== 模型管理接口 =====

@router.get("/models", summary="获取可用模型列表")
async def get_models():
    """获取所有可用的 LLM 模型列表及当前使用的模型"""
    return {
        "models": AVAILABLE_MODELS,
        "current": get_current_model(),
    }


@router.post("/models/set", summary="切换模型")
async def set_model(req: SetModelRequest):
    """切换当前使用的 LLM 模型"""
    success = set_current_model(req.model_id)
    if success:
        return {"success": True, "message": f"已切换到模型: {req.model_id}", "current": req.model_id}
    else:
        valid_ids = [m["id"] for m in AVAILABLE_MODELS]
        raise HTTPException(status_code=400, detail=f"不支持的模型: {req.model_id}，可选: {valid_ids}")
