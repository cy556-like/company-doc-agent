"""
一键补丁脚本：ECS更新 - 添加模型选择 + GLM-5系列

使用方法：
1. 将此文件放到 C:\company-doc-agent\ 目录
2. 运行: python patch_v3.py
3. 重启服务: taskkill /F /IM python.exe && python -m app.main
"""
import os, sys

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

def write_file(rel_path, content):
    full_path = os.path.join(BASE_DIR, rel_path)
    os.makedirs(os.path.dirname(full_path), exist_ok=True)
    with open(full_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"  OK {rel_path}")

# ========== 1. config.py ==========
def patch_config():
    write_file("app/config.py", '''\
"""
应用配置管理
支持动态切换 LLM 模型
"""
import os
from dotenv import load_dotenv

load_dotenv()

# 可用的 LLM 模型列表
AVAILABLE_MODELS = [
    # GLM-5 系列（最新）
    {"id": "glm-5.1", "name": "GLM-5.1", "desc": "最新旗舰，Coding对齐Claude Opus 4.6"},
    {"id": "glm-5-turbo", "name": "GLM-5-Turbo", "desc": "高智能基座，Agent能力SOTA"},
    {"id": "glm-5", "name": "GLM-5", "desc": "高智能基座，编程对齐Claude Opus 4.5"},
    # GLM-4.7 系列
    {"id": "glm-4.7", "name": "GLM-4.7", "desc": "高性能，综合能力提升"},
    {"id": "glm-4.7-flash", "name": "GLM-4.7-Flash", "desc": "快速版，性价比高"},
    # GLM-4 系列（经典）
    {"id": "glm-4-plus", "name": "GLM-4-Plus", "desc": "高性能，复杂任务首选"},
    {"id": "glm-4-long", "name": "GLM-4-Long", "desc": "超长上下文，支持128K"},
    {"id": "glm-4-flash", "name": "GLM-4-Flash", "desc": "最快，适合日常对话"},
    {"id": "glm-4-air", "name": "GLM-4-Air", "desc": "均衡，速度与质量兼顾"},
    {"id": "glm-4-air-0111", "name": "GLM-4-Air-0111", "desc": "Air升级版，效果更好"},
    {"id": "glm-4", "name": "GLM-4", "desc": "经典旗舰模型"},
]


class Settings:
    """应用配置"""
    LLM_API_KEY: str = os.getenv("LLM_API_KEY", "")
    LLM_BASE_URL: str = os.getenv("LLM_BASE_URL", "https://open.bigmodel.cn/api/paas/v4")
    LLM_MODEL: str = os.getenv("LLM_MODEL", "glm-4-flash")
    EMBEDDING_MODEL: str = os.getenv("EMBEDDING_MODEL", "embedding-3")
    APP_HOST: str = os.getenv("APP_HOST", "0.0.0.0")
    APP_PORT: int = int(os.getenv("APP_PORT", "8000"))
    DATA_DIR: str = os.getenv("DATA_DIR", os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data"))
    DOCUMENTS_DIR: str = os.getenv("DOCUMENTS_DIR", os.path.join(DATA_DIR, "documents"))
    CHROMA_DIR: str = os.getenv("CHROMA_DIR", os.path.join(DATA_DIR, "chroma_db"))
    EMPLOYEES_FILE: str = os.getenv("EMPLOYEES_FILE", os.path.join(DATA_DIR, "employees.json"))


settings = Settings()


def set_current_model(model_id: str) -> bool:
    valid_ids = [m["id"] for m in AVAILABLE_MODELS]
    if model_id in valid_ids:
        settings.LLM_MODEL = model_id
        from app.agent.core import reset_agent
        reset_agent()
        return True
    return False


def get_current_model() -> str:
    return settings.LLM_MODEL
''')

# ========== 2. routes.py ==========
def patch_routes():
    write_file("app/api/routes.py", '''\
"""
FastAPI 路由定义
"""
import os, shutil
from typing import Optional
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Query
from pydantic import BaseModel
from app.agent.core import chat
from app.rag.document import index_document, search_documents, list_indexed_documents, read_document_content
from app.memory.manager import (
    get_history_messages, clear_session_history,
    create_chat, list_chats, delete_chat, rename_chat, update_chat_time,
)
from app.auth.user_manager import login_user, register_user
from app.config import settings, AVAILABLE_MODELS, set_current_model, get_current_model

router = APIRouter()


class ChatRequest(BaseModel):
    message: str
    session_id: str = "default"

class ChatResponse(BaseModel):
    response: str
    session_id: str

class SearchRequest(BaseModel):
    query: str
    top_k: int = 3

class AuthRequest(BaseModel):
    username: str
    password: str

class RenameChatRequest(BaseModel):
    username: str
    chat_id: str
    new_title: str

class SetModelRequest(BaseModel):
    model_id: str


# ===== 认证 =====
@router.post("/auth/login")
async def auth_login(req: AuthRequest):
    return login_user(req.username, req.password)

@router.post("/auth/register")
async def auth_register(req: AuthRequest):
    return register_user(req.username, req.password)


# ===== 会话 =====
@router.post("/chats")
async def create_new_chat(username: str = Query(...), title: str = Query("新对话")):
    return {"success": True, "chat": create_chat(username, title)}

@router.get("/chats")
async def get_user_chats(username: str = Query(...)):
    return {"success": True, "chats": list_chats(username)}

@router.delete("/chats/{chat_id}")
async def delete_user_chat(chat_id: str, username: str = Query(...)):
    delete_chat(username, chat_id)
    return {"success": True, "message": "会话已删除"}

@router.put("/chats/{chat_id}/rename")
async def rename_user_chat(chat_id: str, req: RenameChatRequest):
    rename_chat(req.username, req.chat_id, req.new_title)
    return {"success": True, "message": "会话已重命名"}


# ===== Agent 对话 =====
@router.post("/chat", response_model=ChatResponse)
async def chat_api(req: ChatRequest):
    try:
        response = chat(req.message, req.session_id)
        parts = req.session_id.rsplit("_", 1)
        if len(parts) == 2:
            try: update_chat_time(parts[0], req.session_id)
            except: pass
        return ChatResponse(response=response, session_id=req.session_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Agent 处理失败: {str(e)}")


@router.post("/chat-with-file")
async def chat_with_file(
    file: UploadFile = File(...),
    message: str = Form(...),
    session_id: str = Form("default"),
):
    allowed_ext = {".pdf", ".txt", ".docx"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")

    file_path = os.path.join(settings.DOCUMENTS_DIR, file.filename)
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        try: index_document(file_path, file.filename)
        except: pass

        try: file_content = read_document_content(file_path)
        except Exception as e: file_content = f"(无法读取文件内容: {str(e)})"

        max_chars = 8000
        if len(file_content) > max_chars:
            file_content = file_content[:max_chars] + f"\\n...(已截断，共{len(file_content)}字符)"

        enhanced_message = f"""[用户上传了文件: {file.filename}]

文件内容如下：
---
{file_content}
---

文件保存路径: {file_path}

用户的问题/要求: {message}"""

        response = chat(enhanced_message, session_id)
        parts = session_id.rsplit("_", 1)
        if len(parts) == 2:
            try: update_chat_time(parts[0], session_id)
            except: pass
        return {"response": response, "session_id": session_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"处理失败: {str(e)}")


# ===== 文档管理 =====
@router.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    allowed_ext = {".pdf", ".txt", ".docx"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")
    file_path = os.path.join(settings.DOCUMENTS_DIR, file.filename)
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
    try:
        result = index_document(file_path, file.filename)
        return {"status": "success", "detail": result}
    except Exception as e:
        os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"文档索引失败: {str(e)}")


@router.post("/modify-document")
async def modify_document(
    file: UploadFile = File(...),
    instruction: str = Form(...),
    username: str = Form("default"),
):
    allowed_ext = {".pdf", ".txt", ".docx"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")
    temp_dir = os.path.join(settings.DATA_DIR, "temp")
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, file.filename)
    with open(temp_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
    try:
        content = read_document_content(temp_path)
        from langchain_openai import ChatOpenAI
        from langchain_core.messages import SystemMessage, HumanMessage
        llm = ChatOpenAI(api_key=settings.LLM_API_KEY, base_url=settings.LLM_BASE_URL, model=settings.LLM_MODEL, temperature=0.3)
        system_prompt = "你是一个文档修改助手。按修改要求修改文档，只返回修改后的完整内容，不要解释。保持原文档格式和结构，用中文输出。"
        messages = [SystemMessage(content=system_prompt), HumanMessage(content=f"原始文档内容：\\n\\n{content}\\n\\n修改要求：{instruction}")]
        response = llm.invoke(messages)
        modified_content = response.content
        static_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static")
        modified_dir = os.path.join(static_dir, "modified")
        os.makedirs(modified_dir, exist_ok=True)
        output_filename = f"modified_{file.filename}"
        output_path = os.path.join(modified_dir, output_filename)
        if ext == ".pdf":
            from app.utils.pdf_generator import generate_pdf
            success, actual_path = generate_pdf(modified_content, output_path)
            output_filename = os.path.basename(actual_path)
        elif ext == ".docx":
            try:
                from docx import Document
                doc = Document(temp_path)
                for p in doc.paragraphs: p.text = ""
                paras = modified_content.split("\\n")
                if doc.paragraphs: doc.paragraphs[0].text = paras[0] if paras else ""
                for pt in paras[1:]: doc.add_paragraph(pt)
                doc.save(output_path)
            except ImportError:
                output_path = output_path.replace(".docx", ".txt")
                output_filename = os.path.basename(output_path)
                with open(output_path, "w", encoding="utf-8") as f: f.write(modified_content)
        else:
            with open(output_path, "w", encoding="utf-8") as f: f.write(modified_content)
        os.remove(temp_path)
        return {"success": True, "message": "文档修改完成！", "download_url": f"/static/modified/{output_filename}", "filename": output_filename}
    except Exception as e:
        if os.path.exists(temp_path): os.remove(temp_path)
        raise HTTPException(status_code=500, detail=f"文档修改失败: {str(e)}")


@router.post("/search")
async def search_api(req: SearchRequest):
    return {"query": req.query, "results": search_documents(req.query, req.top_k)}

@router.get("/documents")
async def list_documents():
    docs = list_indexed_documents()
    return {"documents": docs, "count": len(docs)}


# ===== 历史 =====
@router.get("/history/{session_id}")
async def get_history(session_id: str):
    msgs = get_history_messages(session_id)
    return {"session_id": session_id, "messages": msgs, "count": len(msgs)}

@router.delete("/history/{session_id}")
async def delete_history(session_id: str):
    clear_session_history(session_id)
    return {"status": "success", "message": f"会话 {session_id} 的历史已清除"}


# ===== 模型管理 =====
@router.get("/models")
async def get_models():
    return {"models": AVAILABLE_MODELS, "current": get_current_model()}

@router.post("/models/set")
async def set_model(req: SetModelRequest):
    success = set_current_model(req.model_id)
    if success:
        return {"success": True, "message": f"已切换到模型: {req.model_id}", "current": req.model_id}
    else:
        valid_ids = [m["id"] for m in AVAILABLE_MODELS]
        raise HTTPException(status_code=400, detail=f"不支持的模型: {req.model_id}，可选: {valid_ids}")
''')

# ========== 3. agent/core.py ==========
def patch_agent_core():
    write_file("app/agent/core.py", '''\
"""
Agent 核心逻辑 - LangGraph ReAct 模式
"""
from typing import Annotated
from typing_extensions import TypedDict
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage, ToolMessage
from langgraph.graph import StateGraph, END
from langgraph.graph.message import add_messages
from langgraph.prebuilt import ToolNode, tools_condition
from app.config import settings
from app.agent.tools import ALL_TOOLS
from app.agent.prompts import SYSTEM_PROMPT
from app.memory.manager import get_session_history

MAX_HISTORY_MESSAGES = 10
MAX_TOOL_ROUNDS = 3


class AgentState(TypedDict):
    messages: Annotated[list, add_messages]


def create_llm():
    return ChatOpenAI(
        api_key=settings.LLM_API_KEY,
        base_url=settings.LLM_BASE_URL,
        model=settings.LLM_MODEL,
        temperature=0.1,
    )


def create_agent_graph():
    llm = create_llm()
    llm_with_tools = llm.bind_tools(ALL_TOOLS)

    def think(state: AgentState):
        messages = state["messages"]
        system_msg = SystemMessage(content=SYSTEM_PROMPT)
        response = llm_with_tools.invoke([system_msg] + messages)
        return {"messages": [response]}

    tool_node = ToolNode(ALL_TOOLS)

    def should_continue(state: AgentState):
        messages = state["messages"]
        tool_message_count = sum(1 for m in messages if isinstance(m, ToolMessage))
        if tool_message_count >= MAX_TOOL_ROUNDS:
            return END
        last_message = messages[-1]
        if hasattr(last_message, "tool_calls") and last_message.tool_calls:
            return "act"
        return END

    graph = StateGraph(AgentState)
    graph.add_node("think", think)
    graph.add_node("act", tool_node)
    graph.set_entry_point("think")
    graph.add_conditional_edges("think", should_continue, {"act": "act", END: END})
    graph.add_edge("act", "think")
    return graph.compile()


_agent_graph = None


def get_agent():
    global _agent_graph
    if _agent_graph is None:
        _agent_graph = create_agent_graph()
    return _agent_graph


def reset_agent():
    global _agent_graph
    _agent_graph = None


def chat(user_input: str, session_id: str = "default") -> str:
    agent = get_agent()
    history = get_session_history(session_id)
    recent_messages = history.messages[-MAX_HISTORY_MESSAGES:]
    all_messages = recent_messages + [HumanMessage(content=user_input)]
    result = agent.invoke({"messages": all_messages})
    ai_message = result["messages"][-1]
    history.add_message(HumanMessage(content=user_input))
    history.add_message(ai_message)
    return ai_message.content
''')


# ========== main ==========
def main():
    print("=" * 55)
    print("  ECS 补丁 v3 - 模型选择 + GLM-5系列")
    print("=" * 55)
    print()
    patch_config()
    patch_routes()
    patch_agent_core()
    print()
    print("Done! Updated files:")
    print("  app/config.py      - 11 models (GLM-5/5.1/5-Turbo/4.7)")
    print("  app/api/routes.py  - /models + /models/set API")
    print("  app/agent/core.py  - reset_agent() hot-swap")
    print()
    print("NOTE: index.html not changed (if model select works, skip)")
    print("      If UI has no model dropdown, also run: patch_html.py")
    print()
    print("Restart: taskkill /F /IM python.exe && python -m app.main")


if __name__ == "__main__":
    main()
