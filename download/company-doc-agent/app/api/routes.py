"""
FastAPI 路由定义
提供 REST API 接口供前端和外部调用
"""
import os
import shutil
import tempfile
from typing import Optional

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Query
from pydantic import BaseModel

from app.agent.core import chat
from app.rag.document import index_document, search_documents, list_indexed_documents
from app.memory.manager import (
    get_history_messages, clear_session_history,
    create_chat, list_chats, delete_chat, rename_chat, update_chat_time,
)
from app.auth.user_manager import login_user, register_user
from app.config import settings

router = APIRouter()


# ===== 请求/响应模型 =====
class ChatRequest(BaseModel):
    """聊天请求"""
    message: str
    session_id: str = "default"


class ChatResponse(BaseModel):
    """聊天响应"""
    response: str
    session_id: str


class SearchRequest(BaseModel):
    """文档搜索请求"""
    query: str
    top_k: int = 3


class AuthRequest(BaseModel):
    """认证请求"""
    username: str
    password: str


class RenameChatRequest(BaseModel):
    """重命名会话请求"""
    username: str
    chat_id: str
    new_title: str


# ===== 认证接口 =====

@router.post("/auth/login", summary="用户登录")
async def auth_login(req: AuthRequest):
    """用户登录验证"""
    result = login_user(req.username, req.password)
    return result


@router.post("/auth/register", summary="用户注册")
async def auth_register(req: AuthRequest):
    """用户注册"""
    result = register_user(req.username, req.password)
    return result


# ===== 会话管理接口 =====

@router.post("/chats", summary="创建新会话")
async def create_new_chat(username: str = Query(..., description="用户名"), title: str = Query("新对话", description="会话标题")):
    """为用户创建一个新的聊天会话"""
    chat_info = create_chat(username, title)
    return {"success": True, "chat": chat_info}


@router.get("/chats", summary="获取用户的所有会话")
async def get_user_chats(username: str = Query(..., description="用户名")):
    """获取用户的所有聊天会话列表"""
    chats = list_chats(username)
    return {"success": True, "chats": chats}


@router.delete("/chats/{chat_id}", summary="删除会话")
async def delete_user_chat(chat_id: str, username: str = Query(..., description="用户名")):
    """删除指定的聊天会话"""
    delete_chat(username, chat_id)
    return {"success": True, "message": "会话已删除"}


@router.put("/chats/{chat_id}/rename", summary="重命名会话")
async def rename_user_chat(chat_id: str, req: RenameChatRequest):
    """重命名指定的聊天会话"""
    rename_chat(req.username, req.chat_id, req.new_title)
    return {"success": True, "message": "会话已重命名"}


# ===== 核心 Agent 接口 =====

@router.post("/chat", response_model=ChatResponse, summary="与 Agent 对话")
async def chat_api(req: ChatRequest):
    """
    核心接口：与文档助手 Agent 对话

    - 支持 RAG 文档问答
    - 支持员工信息查询
    - 支持多轮对话
    """
    try:
        response = chat(req.message, req.session_id)
        # 更新会话时间（从 session_id 中提取 username）
        # session_id 格式: username_xxxxx
        parts = req.session_id.rsplit("_", 1)
        if len(parts) == 2:
            username = parts[0]
            try:
                update_chat_time(username, req.session_id)
            except Exception:
                pass
        return ChatResponse(response=response, session_id=req.session_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Agent 处理失败: {str(e)}")


# ===== 文档管理接口 =====

@router.post("/upload", summary="上传文档到知识库")
async def upload_document(file: UploadFile = File(...)):
    """
    上传文档并自动索引到向量数据库
    支持 PDF、TXT、DOCX 格式
    """
    # 检查文件格式
    allowed_ext = {".pdf", ".txt", ".docx"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: {ext}，仅支持 {allowed_ext}",
        )

    # 保存文件
    file_path = os.path.join(settings.DOCUMENTS_DIR, file.filename)
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # 索引文档
    try:
        result = index_document(file_path, file.filename)
        return {"status": "success", "detail": result}
    except Exception as e:
        # 索引失败则删除文件
        os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"文档索引失败: {str(e)}")


@router.post("/modify-document", summary="修改文档")
async def modify_document(
    file: UploadFile = File(...),
    instruction: str = Form(...),
    username: str = Form("default"),
):
    """
    上传文档并根据修改要求进行修改，返回修改后的文件下载链接

    - 支持 PDF、TXT、DOCX 格式
    - 修改要求：用自然语言描述如何修改
    - 返回修改后的文件下载链接
    """
    # 检查文件格式
    allowed_ext = {".pdf", ".txt", ".docx"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: {ext}，仅支持 {allowed_ext}",
        )

    # 保存上传的文件到临时目录
    temp_dir = os.path.join(settings.DATA_DIR, "temp")
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, file.filename)
    with open(temp_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        # 读取文档内容
        from app.rag.document import read_document_content
        content = read_document_content(temp_path)

        # 调用 LLM 修改文档
        from langchain_openai import ChatOpenAI
        from langchain_core.messages import SystemMessage, HumanMessage

        llm = ChatOpenAI(
            api_key=settings.LLM_API_KEY,
            base_url=settings.LLM_BASE_URL,
            model=settings.LLM_MODEL,
            temperature=0.3,
        )

        system_prompt = f"""你是一个文档修改助手。用户会给你一份文档的原始内容和修改要求，你需要按照修改要求对文档进行修改，然后返回修改后的完整文档内容。

规则：
1. 只返回修改后的文档内容，不要添加任何解释说明
2. 保持原文档的格式和结构
3. 只修改用户要求的部分，其余内容保持不变
4. 如果是表格数据，保持表格格式
5. 用中文输出"""

        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=f"原始文档内容：\n\n{content}\n\n修改要求：{instruction}"),
        ]

        response = llm.invoke(messages)
        modified_content = response.content

        # 保存修改后的文件（保存到 static/modified 目录，可通过 /static/modified/ 下载）
        static_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static")
        modified_dir = os.path.join(static_dir, "modified")
        os.makedirs(modified_dir, exist_ok=True)

        output_filename = f"modified_{file.filename}"
        output_path = os.path.join(modified_dir, output_filename)

        if ext == ".docx":
            # 使用 python-docx 保存 docx 格式
            try:
                from docx import Document
                doc = Document(temp_path)
                # 清空所有段落
                for paragraph in doc.paragraphs:
                    paragraph.text = ""
                # 写入修改后的内容
                paragraphs = modified_content.split("\n")
                # 第一个段落使用已有的第一个段落
                if doc.paragraphs:
                    doc.paragraphs[0].text = paragraphs[0] if paragraphs else ""
                # 剩余段落添加
                for p_text in paragraphs[1:]:
                    doc.add_paragraph(p_text)
                doc.save(output_path)
            except ImportError:
                # 如果没有 python-docx，就保存为 txt
                output_filename = f"modified_{os.path.splitext(file.filename)[0]}.txt"
                output_path = os.path.join(modified_dir, output_filename)
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(modified_content)
        elif ext == ".pdf":
            # 使用 fpdf2 生成真正的 PDF 文件（支持中文）
            from app.utils.pdf_generator import generate_pdf
            success, actual_path = generate_pdf(modified_content, output_path, title=f"修改后的 {file.filename}")
            # 更新实际输出文件名（可能回退为 .txt）
            output_filename = os.path.basename(actual_path)
        else:
            # .txt 等其他格式直接保存文本
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(modified_content)

        # 清理临时文件
        os.remove(temp_path)

        download_url = f"/static/modified/{output_filename}"

        return {
            "success": True,
            "message": f"文档修改完成！已按您的修改要求处理。",
            "download_url": download_url,
            "filename": output_filename,
        }

    except Exception as e:
        # 清理临时文件
        if os.path.exists(temp_path):
            os.remove(temp_path)
        raise HTTPException(status_code=500, detail=f"文档修改失败: {str(e)}")


@router.post("/search", summary="搜索文档内容")
async def search_api(req: SearchRequest):
    """在文档库中搜索相关内容"""
    results = search_documents(req.query, req.top_k)
    return {"query": req.query, "results": results}


@router.get("/documents", summary="列出所有已索引文档")
async def list_documents():
    """获取知识库中所有文档列表"""
    docs = list_indexed_documents()
    return {"documents": docs, "count": len(docs)}


# ===== 对话历史接口 =====

@router.get("/history/{session_id}", summary="获取对话历史")
async def get_history(session_id: str):
    """获取指定会话的对话历史"""
    messages = get_history_messages(session_id)
    return {"session_id": session_id, "messages": messages, "count": len(messages)}


@router.delete("/history/{session_id}", summary="清除对话历史")
async def delete_history(session_id: str):
    """清除指定会话的对话历史"""
    clear_session_history(session_id)
    return {"status": "success", "message": f"会话 {session_id} 的历史已清除"}
