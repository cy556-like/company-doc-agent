"""
Agent 工具定义模块
每个工具 = Agent 的一个「能力」
Agent 会根据用户问题自动选择调用哪个工具
"""
import json
import os
from typing import Optional

from langchain_core.tools import tool

from app.config import settings
from app.rag.document import search_documents, index_document, list_indexed_documents


@tool
def search_documents_tool(query: str) -> str:
    """在公司文档库中搜索与查询相关的文档内容。当用户问关于公司制度、流程、规范等问题时使用。

    Args:
        query: 搜索查询内容
    """
    results = search_documents(query, top_k=3)

    if not results:
        return "未找到相关文档内容。"

    output = "检索到以下相关内容：\n\n"
    for i, r in enumerate(results, 1):
        output += f"【文档 {i}】来源: {r['source']}\n"
        output += f"{r['content']}\n"
        output += f"(相似度: {r['relevance_score']})\n\n"

    return output


@tool
def lookup_employee_tool(name: str = "", department: str = "") -> str:
    """查询公司员工信息，包括姓名、部门、职位、邮箱等。当用户问「某某在哪个部门」、「某部门有哪些人」等问题时使用。

    Args:
        name: 员工姓名（可选，模糊匹配）
        department: 部门名称（可选）
    """
    employees_file = settings.EMPLOYEES_FILE

    if not os.path.exists(employees_file):
        return "员工数据库暂未初始化，请先运行 scripts/seed_data.py 初始化数据。"

    with open(employees_file, "r", encoding="utf-8") as f:
        employees = json.load(f)

    results = employees

    # 按姓名过滤
    if name:
        results = [e for e in results if name in e.get("name", "")]

    # 按部门过滤
    if department:
        results = [e for e in results if department in e.get("department", "")]

    if not results:
        return f"未找到匹配的员工信息。（搜索条件：姓名={name}, 部门={department}）"

    output = f"找到 {len(results)} 位员工：\n\n"
    for e in results:
        output += f"{e['name']} | 部门: {e['department']} | 职位: {e['position']} | 邮箱: {e['email']}\n"
        if e.get("phone"):
            output += f"   电话: {e['phone']}\n"

    return output


@tool
def list_documents_tool() -> str:
    """列出知识库中所有可用的文档。当用户想了解有哪些文档时使用。"""
    docs = list_indexed_documents()

    if not docs:
        return "知识库中暂无文档。请先上传文档。"

    output = f"知识库中共有 {len(docs)} 个文档：\n\n"
    for i, doc in enumerate(docs, 1):
        output += f"  {i}. {doc}\n"

    return output


@tool
def upload_document_tool(file_path: str) -> str:
    """上传新文档到知识库。当用户需要添加新文档时使用。

    Args:
        file_path: 文档文件路径
    """
    if not os.path.exists(file_path):
        return f"文件不存在: {file_path}"

    try:
        result = index_document(file_path)
        return f"文档上传成功！{result['message']}"
    except Exception as e:
        return f"文档上传失败: {str(e)}"


@tool
def modify_document_tool(file_path: str, instruction: str) -> str:
    """修改已有文档的内容。当用户需要修改、编辑、更新文档时使用。

    Args:
        file_path: 要修改的文档路径
        instruction: 修改要求（自然语言描述如何修改）
    """
    if not os.path.exists(file_path):
        return f"文件不存在: {file_path}"

    try:
        from app.rag.document import read_document_content
        from langchain_openai import ChatOpenAI
        from langchain_core.messages import SystemMessage, HumanMessage

        # 读取文档内容
        content = read_document_content(file_path)

        # 调用 LLM 修改文档
        llm = ChatOpenAI(
            api_key=settings.LLM_API_KEY,
            base_url=settings.LLM_BASE_URL,
            model=settings.LLM_MODEL,
            temperature=0.3,
        )

        system_prompt = """你是一个文档修改助手。用户会给你一份文档的原始内容和修改要求，你需要按照修改要求对文档进行修改，然后返回修改后的完整文档内容。

规则：
1. 只返回修改后的文档内容，不要添加任何解释说明
2. 保持原文档的格式和结构
3. 只修改用户要求的部分，其余内容保持不变
4. 用中文输出"""

        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=f"原始文档内容：\n\n{content}\n\n修改要求：{instruction}"),
        ]

        response = llm.invoke(messages)
        modified_content = response.content

        # 保存修改后的文件（保存到 static/modified 目录，可通过 /static/modified/ 下载）
        ext = os.path.splitext(file_path)[1].lower()
        static_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static")
        modified_dir = os.path.join(static_dir, "modified")
        os.makedirs(modified_dir, exist_ok=True)
        output_filename = f"modified_{os.path.basename(file_path)}"
        output_path = os.path.join(modified_dir, output_filename)

        if ext == ".docx":
            try:
                from docx import Document
                doc = Document(file_path)
                for paragraph in doc.paragraphs:
                    paragraph.text = ""
                paragraphs = modified_content.split("\n")
                if doc.paragraphs:
                    doc.paragraphs[0].text = paragraphs[0] if paragraphs else ""
                for p_text in paragraphs[1:]:
                    doc.add_paragraph(p_text)
                doc.save(output_path)
            except ImportError:
                output_path = output_path.replace(".docx", ".txt")
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(modified_content)
        else:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(modified_content)

        return f"文档修改完成！修改后的文件已保存，下载链接: /static/modified/{output_filename}"

    except Exception as e:
        return f"文档修改失败: {str(e)}"


# ===== 导出所有工具列表 =====
ALL_TOOLS = [
    search_documents_tool,
    lookup_employee_tool,
    list_documents_tool,
    upload_document_tool,
    modify_document_tool,
]
