"""
一键补丁脚本 v2 - 修复所有问题
修复内容：
1. ChatGPT风格文件对话 - 文件+消息走Agent自动判断意图，默认文字回答
2. 延迟导入ChromaDB - 即使ChromaDB DLL报错也不影响文档读取和对话
3. PDF生成使用fpdf2 - 不再生成损坏的PDF文件
4. 前端UI - 📎文件标记，动态placeholder

运行：python patch_v2.py
"""
import os

BASE = r"C:\company-doc-agent"

def write_file(rel_path, content):
    """写入文件，自动创建目录"""
    full_path = os.path.join(BASE, rel_path)
    os.makedirs(os.path.dirname(full_path), exist_ok=True)
    with open(full_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"  OK {rel_path}")

print("=" * 50)
print("开始补丁 v2 ...")
print("=" * 50)

# ========== 1. app/utils/__init__.py ==========
write_file(r"app\utils\__init__.py", "")

# ========== 2. app/utils/pdf_generator.py ==========
write_file(r"app\utils\pdf_generator.py", r'''"""
PDF 生成工具模块
使用 fpdf2 生成 PDF，支持中文字体
如果 fpdf2 不可用，自动降级为保存 .txt 文件
"""
import os
import platform
import logging

logger = logging.getLogger(__name__)

CHINESE_FONT_PATHS = {
    "Windows": [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/msyh.ttf",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simhei.ttf",
    ],
    "Linux": [
        "/usr/share/fonts/truetype/chinese/NotoSansSC-Regular.ttf",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
    ],
    "Darwin": [
        "/System/Library/Fonts/PingFang.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ],
}


def find_chinese_font():
    """查找系统中可用的中文字体"""
    system = platform.system()
    for path in CHINESE_FONT_PATHS.get(system, []):
        if os.path.exists(path):
            return path
    return None


def generate_pdf(text, output_path, title="修改后的文档"):
    """
    生成 PDF 文件

    Args:
        text: 文本内容
        output_path: 输出路径
        title: 文档标题

    Returns:
        tuple: (success: bool, actual_path: str)
    """
    try:
        from fpdf import FPDF
    except ImportError:
        logger.warning("fpdf2 未安装，降级为 txt 文件")
        return _save_as_txt(text, output_path)

    font_path = find_chinese_font()
    if not font_path:
        logger.warning("未找到中文字体，降级为 txt 文件")
        return _save_as_txt(text, output_path)

    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_font("ChineseFont", "", font_path, uni=True)
        pdf.set_font("ChineseFont", "", 12)

        for line in text.split("\n"):
            if not line.strip():
                pdf.ln(6)
                continue
            pdf.multi_cell(0, 7, line)

        pdf.output(output_path)
        return True, output_path

    except Exception as e:
        logger.error(f"PDF 生成失败: {e}，降级为 txt 文件")
        return _save_as_txt(text, output_path)


def _save_as_txt(text, output_path):
    """降级方案：保存为 txt 文件"""
    if output_path.lower().endswith(".pdf"):
        output_path = output_path[:-4] + ".txt"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)
    return True, output_path
''')

# ========== 3. app/rag/document.py - 延迟导入ChromaDB ==========
write_file(r"app\rag\document.py", r'''"""
文档处理与向量化模块 (RAG)
修复：延迟导入 ChromaDB，即使 DLL 报错也不影响文档读取
"""
import os
from typing import Optional

from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import settings

# ===== 单例缓存 =====
_vector_store = None
_embeddings = None
_chromadb_available = True


def _check_chromadb():
    """检查 ChromaDB 是否可用"""
    global _chromadb_available
    if not _chromadb_available:
        return False
    try:
        from langchain_chroma import Chroma
        return True
    except ImportError:
        try:
            from langchain_community.vectorstores import Chroma
            return True
        except Exception:
            _chromadb_available = False
            return False
    except Exception:
        _chromadb_available = False
        return False


def get_embeddings():
    """获取 Embedding 模型（单例缓存）"""
    global _embeddings
    if _embeddings is None:
        from langchain_openai import OpenAIEmbeddings
        embedding_model = getattr(settings, 'EMBEDDING_MODEL', 'embedding-3')
        _embeddings = OpenAIEmbeddings(
            api_key=settings.LLM_API_KEY,
            base_url=settings.LLM_BASE_URL,
            model=embedding_model,
        )
    return _embeddings


def get_vector_store():
    """获取 ChromaDB 向量数据库（延迟导入，失败不崩溃）"""
    global _vector_store, _chromadb_available
    if _vector_store is not None:
        return _vector_store
    if not _chromadb_available:
        return None
    try:
        try:
            from langchain_chroma import Chroma
        except ImportError:
            from langchain_community.vectorstores import Chroma
        embeddings = get_embeddings()
        _vector_store = Chroma(
            persist_directory=settings.CHROMA_DIR,
            embedding_function=embeddings,
        )
        return _vector_store
    except Exception as e:
        print(f"[WARN] ChromaDB 不可用: {e}")
        _chromadb_available = False
        return None


def load_document(file_path: str) -> list:
    """根据文件类型加载文档，不依赖 ChromaDB"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
    elif ext == ".txt":
        loader = TextLoader(file_path, encoding="utf-8")
    elif ext == ".docx":
        loader = Docx2txtLoader(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")
    return loader.load()


def read_document_content(file_path: str) -> str:
    """读取文档纯文本内容，不依赖 ChromaDB"""
    docs = load_document(file_path)
    content = "\n\n".join([doc.page_content for doc in docs])
    return content


def split_documents(docs: list, chunk_size: int = 500, chunk_overlap: int = 100) -> list:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", "；", "，", " ", ""],
    )
    return splitter.split_documents(docs)


def index_document(file_path: str, filename: str = None) -> dict:
    """索引文档到向量库，ChromaDB 不可用时只做分块不存储"""
    if filename is None:
        filename = os.path.basename(file_path)

    docs = load_document(file_path)
    for doc in docs:
        doc.metadata["source_file"] = filename
    chunks = split_documents(docs)

    vector_store = get_vector_store()
    if vector_store is not None:
        vector_store.add_documents(chunks)
        return {
            "filename": filename,
            "chunks": len(chunks),
            "status": "success",
            "message": f"文档 {filename} 已成功索引，共 {len(chunks)} 个分块",
        }
    else:
        return {
            "filename": filename,
            "chunks": len(chunks),
            "status": "partial",
            "message": f"文档 {filename} 已读取（{len(chunks)} 个分块），但 ChromaDB 不可用，暂未存入向量库",
        }


def search_documents(query: str, top_k: int = 3) -> list[dict]:
    """搜索文档，ChromaDB 不可用时返回空"""
    vector_store = get_vector_store()
    if vector_store is None:
        return []
    results = vector_store.similarity_search_with_score(query, k=top_k)
    formatted = []
    for doc, score in results:
        formatted.append({
            "content": doc.page_content,
            "source": doc.metadata.get("source_file", "未知来源"),
            "relevance_score": round(1 - score, 4),
        })
    return formatted


def list_indexed_documents() -> list[str]:
    """列出已索引文档，ChromaDB 不可用时返回空"""
    vector_store = get_vector_store()
    if vector_store is None:
        return []
    try:
        collection = vector_store._collection
        all_docs = collection.get(include=["metadatas"])
        sources = set()
        for meta in all_docs["metadatas"]:
            if meta and "source_file" in meta:
                sources.add(meta["source_file"])
        return sorted(list(sources))
    except Exception:
        return []
''')

# ========== 4. app/agent/prompts.py ==========
write_file(r"app\agent\prompts.py", r'''"""
Prompt 模板模块
定义 Agent 的系统提示词和各工具的描述
"""

SYSTEM_PROMPT = """你是一个企业文档智能助手，名叫「小智」。你服务于公司内部员工，帮助他们高效地获取和处理公司文档信息。

## 你的核心能力
1. **文档问答**：根据公司文档回答员工的问题（基于 RAG 检索）
2. **员工查询**：查询公司员工的基本信息、部门、职位等
3. **文档搜索**：搜索包含特定关键词的文档
4. **文档上传**：帮助员工上传新的文档到知识库
5. **文档修改**：根据用户要求修改已有文档的内容

## 工作原则
- 回答必须基于检索到的文档内容，不要凭空编造
- 如果找不到相关信息，明确告知用户，不要猜测
- 回答时引用来源文档，增加可信度
- 用中文回答，简洁专业

## 回答格式
- 先给出直接回答
- 再补充相关细节
- 如果涉及多个文档，逐一说明来源

## 用户上传文件时的处理规则
当用户上传了文件时，消息中会包含 [用户上传了文件: xxx] 标记和文件内容。
- **默认行为**：基于文件内容直接回答用户的问题，用文字给出总结、分析、搜索结果等，不要返回文件。
- **只有当用户明确要求修改/返回文件时**（如"帮我修改"、"返回修改后的文件"、"给我修改后的完整文件"），才调用 modify_document 工具，使用消息中提供的文件保存路径作为 file_path 参数。
- 如果用户只是提问、总结、分析、搜索、提取信息，直接基于已有的文件内容用文字回答即可，不需要调用任何工具。
"""

TOOL_DESCRIPTIONS = {
    "search_documents": "在公司文档库中搜索与查询相关的文档内容。当用户问关于公司制度、流程、规范等问题时使用。",
    "lookup_employee": "查询公司员工信息，包括姓名、部门、职位、邮箱等。当用户问「某某在哪个部门」、「某部门有哪些人」等问题时使用。",
    "list_documents": "列出知识库中所有可用的文档。当用户想了解有哪些文档时使用。",
    "upload_document": "上传新文档到知识库。当用户需要添加新文档时使用。",
    "modify_document": "修改已有文档的内容。当用户明确要求修改文档并返回文件时使用。",
}
''')

# ========== 5. app/agent/tools.py ==========
write_file(r"app\agent\tools.py", r'''"""
Agent 工具定义模块
"""
import json
import os
from typing import Optional

from langchain_core.tools import tool
from app.config import settings
from app.rag.document import search_documents, index_document, list_indexed_documents


@tool
def search_documents_tool(query: str) -> str:
    """在公司文档库中搜索与查询相关的文档内容。当用户问关于公司制度、流程、规范等问题时使用。

    Args:
        query: 搜索查询内容
    """
    results = search_documents(query, top_k=3)
    if not results:
        return "未找到相关文档内容。知识库可能暂时不可用，请稍后再试。"
    output = "检索到以下相关内容：\n\n"
    for i, r in enumerate(results, 1):
        output += f"【文档 {i}】来源: {r['source']}\n"
        output += f"{r['content']}\n"
        output += f"(相似度: {r['relevance_score']})\n\n"
    return output


@tool
def lookup_employee_tool(name: str = "", department: str = "") -> str:
    """查询公司员工信息，包括姓名、部门、职位、邮箱等。

    Args:
        name: 员工姓名（可选，模糊匹配）
        department: 部门名称（可选）
    """
    employees_file = settings.EMPLOYEES_FILE
    if not os.path.exists(employees_file):
        return "员工数据库暂未初始化。"
    with open(employees_file, "r", encoding="utf-8") as f:
        employees = json.load(f)
    results = employees
    if name:
        results = [e for e in results if name in e.get("name", "")]
    if department:
        results = [e for e in results if department in e.get("department", "")]
    if not results:
        return f"未找到匹配的员工信息。（姓名={name}, 部门={department}）"
    output = f"找到 {len(results)} 位员工：\n\n"
    for e in results:
        output += f"{e['name']} | 部门: {e['department']} | 职位: {e['position']} | 邮箱: {e['email']}\n"
    return output


@tool
def list_documents_tool() -> str:
    """列出知识库中所有可用的文档。当用户想了解有哪些文档时使用。"""
    docs = list_indexed_documents()
    if not docs:
        return "知识库中暂无文档或向量库暂时不可用。"
    output = f"知识库中共有 {len(docs)} 个文档：\n\n"
    for i, doc in enumerate(docs, 1):
        output += f"  {i}. {doc}\n"
    return output


@tool
def upload_document_tool(file_path: str) -> str:
    """上传新文档到知识库。

    Args:
        file_path: 文档文件路径
    """
    if not os.path.exists(file_path):
        return f"文件不存在: {file_path}"
    try:
        result = index_document(file_path)
        return f"文档上传成功！{result['message']}"
    except Exception as e:
        return f"文档上传失败: {str(e)}"


@tool
def modify_document_tool(file_path: str, instruction: str) -> str:
    """修改已有文档的内容。当用户明确要求修改文档并返回修改后的文件时使用。

    Args:
        file_path: 要修改的文档路径
        instruction: 修改要求
    """
    if not os.path.exists(file_path):
        return f"文件不存在: {file_path}"
    try:
        from app.rag.document import read_document_content
        from langchain_openai import ChatOpenAI
        from langchain_core.messages import SystemMessage, HumanMessage

        content = read_document_content(file_path)

        llm = ChatOpenAI(
            api_key=settings.LLM_API_KEY,
            base_url=settings.LLM_BASE_URL,
            model=settings.LLM_MODEL,
            temperature=0.3,
        )

        messages = [
            SystemMessage(content="你是一个文档修改助手。按照修改要求修改文档，返回修改后的完整内容。只返回修改后内容，不添加解释。用中文输出。"),
            HumanMessage(content=f"原始文档内容：\n\n{content}\n\n修改要求：{instruction}"),
        ]

        response = llm.invoke(messages)
        modified_content = response.content

        ext = os.path.splitext(file_path)[1].lower()
        static_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static")
        modified_dir = os.path.join(static_dir, "modified")
        os.makedirs(modified_dir, exist_ok=True)
        output_filename = f"modified_{os.path.basename(file_path)}"
        output_path = os.path.join(modified_dir, output_filename)

        if ext == ".docx":
            try:
                from docx import Document
                doc = Document(file_path)
                for paragraph in doc.paragraphs:
                    paragraph.text = ""
                paragraphs = modified_content.split("\n")
                if doc.paragraphs:
                    doc.paragraphs[0].text = paragraphs[0] if paragraphs else ""
                for p_text in paragraphs[1:]:
                    doc.add_paragraph(p_text)
                doc.save(output_path)
            except ImportError:
                output_path = output_path.replace(".docx", ".txt")
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(modified_content)
        elif ext == ".pdf":
            from app.utils.pdf_generator import generate_pdf
            success, actual_path = generate_pdf(modified_content, output_path)
            output_filename = os.path.basename(actual_path)
        else:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(modified_content)

        return f"文档修改完成！修改后的文件已保存，下载链接: /static/modified/{output_filename}"
    except Exception as e:
        return f"文档修改失败: {str(e)}"


ALL_TOOLS = [
    search_documents_tool,
    lookup_employee_tool,
    list_documents_tool,
    upload_document_tool,
    modify_document_tool,
]
''')

# ========== 6. app/api/routes.py ==========
write_file(r"app\api\routes.py", r'''"""
FastAPI 路由定义
"""
import os
import shutil
from typing import Optional

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Query
from pydantic import BaseModel

from app.agent.core import chat
from app.rag.document import index_document, search_documents, list_indexed_documents, read_document_content
from app.memory.manager import (
    get_history_messages, clear_session_history,
    create_chat, list_chats, delete_chat, rename_chat, update_chat_time,
)
from app.auth.user_manager import login_user, register_user
from app.config import settings

router = APIRouter()


class ChatRequest(BaseModel):
    message: str
    session_id: str = "default"


class ChatResponse(BaseModel):
    response: str
    session_id: str


class SearchRequest(BaseModel):
    query: str
    top_k: int = 3


class AuthRequest(BaseModel):
    username: str
    password: str


class RenameChatRequest(BaseModel):
    username: str
    chat_id: str
    new_title: str


@router.post("/auth/login")
async def auth_login(req: AuthRequest):
    return login_user(req.username, req.password)


@router.post("/auth/register")
async def auth_register(req: AuthRequest):
    return register_user(req.username, req.password)


@router.post("/chats")
async def create_new_chat(username: str = Query(...), title: str = Query("新对话")):
    return {"success": True, "chat": create_chat(username, title)}


@router.get("/chats")
async def get_user_chats(username: str = Query(...)):
    return {"success": True, "chats": list_chats(username)}


@router.delete("/chats/{chat_id}")
async def delete_user_chat(chat_id: str, username: str = Query(...)):
    delete_chat(username, chat_id)
    return {"success": True, "message": "会话已删除"}


@router.put("/chats/{chat_id}/rename")
async def rename_user_chat(chat_id: str, req: RenameChatRequest):
    rename_chat(req.username, req.chat_id, req.new_title)
    return {"success": True, "message": "会话已重命名"}


@router.post("/chat", response_model=ChatResponse)
async def chat_api(req: ChatRequest):
    try:
        response = chat(req.message, req.session_id)
        parts = req.session_id.rsplit("_", 1)
        if len(parts) == 2:
            try:
                update_chat_time(parts[0], req.session_id)
            except Exception:
                pass
        return ChatResponse(response=response, session_id=req.session_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Agent 处理失败: {str(e)}")


@router.post("/chat-with-file")
async def chat_with_file(
    file: UploadFile = File(...),
    message: str = Form(...),
    session_id: str = Form("default"),
):
    """ChatGPT风格：文件+消息，Agent自动判断意图，默认文字回答"""
    allowed_ext = {".pdf", ".txt", ".docx"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")

    # 保存文件
    file_path = os.path.join(settings.DOCUMENTS_DIR, file.filename)
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        # 索引到知识库（失败不影响对话）
        try:
            index_document(file_path, file.filename)
        except Exception:
            pass

        # 提取文件内容（不依赖 ChromaDB）
        try:
            file_content = read_document_content(file_path)
        except Exception as e:
            file_content = f"（无法读取文件内容: {str(e)}）"

        # 限制长度
        max_chars = 8000
        if len(file_content) > max_chars:
            file_content = file_content[:max_chars] + f"\n\n...（已截断，共{len(file_content)}字符）"

        # 构建增强消息
        enhanced_message = f"""[用户上传了文件: {file.filename}]

文件内容如下：
---
{file_content}
---

文件保存路径: {file_path}

用户的问题/要求: {message}"""

        # 调用 Agent
        response = chat(enhanced_message, session_id)

        # 更新会话时间
        parts = session_id.rsplit("_", 1)
        if len(parts) == 2:
            try:
                update_chat_time(parts[0], session_id)
            except Exception:
                pass

        return {"response": response, "session_id": session_id}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"处理失败: {str(e)}")


@router.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    allowed_ext = {".pdf", ".txt", ".docx"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")

    file_path = os.path.join(settings.DOCUMENTS_DIR, file.filename)
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        result = index_document(file_path, file.filename)
        return {"status": "success", "detail": result}
    except Exception as e:
        os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"文档索引失败: {str(e)}")


@router.post("/modify-document")
async def modify_document(
    file: UploadFile = File(...),
    instruction: str = Form(...),
    username: str = Form("default"),
):
    allowed_ext = {".pdf", ".txt", ".docx"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")

    temp_dir = os.path.join(settings.DATA_DIR, "temp")
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, file.filename)
    with open(temp_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        content = read_document_content(temp_path)

        from langchain_openai import ChatOpenAI
        from langchain_core.messages import SystemMessage, HumanMessage

        llm = ChatOpenAI(
            api_key=settings.LLM_API_KEY,
            base_url=settings.LLM_BASE_URL,
            model=settings.LLM_MODEL,
            temperature=0.3,
        )

        system_prompt = """你是一个文档修改助手。用户会给你一份文档的原始内容和修改要求，你需要按照修改要求对文档进行修改，然后返回修改后的完整文档内容。
规则：
1. 只返回修改后的文档内容，不要添加任何解释说明
2. 保持原文档的格式和结构
3. 只修改用户要求的部分，其余内容保持不变
4. 如果是表格数据，保持表格格式
5. 用中文输出"""

        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=f"原始文档内容：\n\n{content}\n\n修改要求：{instruction}"),
        ]

        response = llm.invoke(messages)
        modified_content = response.content

        static_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static")
        modified_dir = os.path.join(static_dir, "modified")
        os.makedirs(modified_dir, exist_ok=True)
        output_filename = f"modified_{file.filename}"
        output_path = os.path.join(modified_dir, output_filename)

        if ext == ".docx":
            try:
                from docx import Document
                doc = Document(temp_path)
                for paragraph in doc.paragraphs:
                    paragraph.text = ""
                paragraphs = modified_content.split("\n")
                if doc.paragraphs:
                    doc.paragraphs[0].text = paragraphs[0] if paragraphs else ""
                for p_text in paragraphs[1:]:
                    doc.add_paragraph(p_text)
                doc.save(output_path)
            except ImportError:
                output_filename = f"modified_{os.path.splitext(file.filename)[0]}.txt"
                output_path = os.path.join(modified_dir, output_filename)
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(modified_content)
        elif ext == ".pdf":
            from app.utils.pdf_generator import generate_pdf
            success, actual_path = generate_pdf(modified_content, output_path)
            output_filename = os.path.basename(actual_path)
        else:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(modified_content)

        os.remove(temp_path)
        download_url = f"/static/modified/{output_filename}"

        return {
            "success": True,
            "message": "文档修改完成！",
            "download_url": download_url,
            "filename": output_filename,
        }
    except Exception as e:
        if os.path.exists(temp_path):
            os.remove(temp_path)
        raise HTTPException(status_code=500, detail=f"文档修改失败: {str(e)}")


@router.post("/search")
async def search_api(req: SearchRequest):
    return {"query": req.query, "results": search_documents(req.query, req.top_k)}


@router.get("/documents")
async def list_documents():
    docs = list_indexed_documents()
    return {"documents": docs, "count": len(docs)}


@router.get("/history/{session_id}")
async def get_history(session_id: str):
    messages = get_history_messages(session_id)
    return {"session_id": session_id, "messages": messages, "count": len(messages)}


@router.delete("/history/{session_id}")
async def delete_history(session_id: str):
    clear_session_history(session_id)
    return {"status": "success", "message": f"会话 {session_id} 的历史已清除"}
''')

# ========== 7. index.html - 修改关键部分 ==========
print("\n修改 index.html ...")
html_path = os.path.join(BASE, r"app\static\index.html")
if not os.path.exists(html_path):
    print("  ERROR: index.html 不存在！")
else:
    with open(html_path, "r", encoding="utf-8") as f:
        html = f.read()

    # 1. 替换显示文本: [修改文档] -> 📎
    html = html.replace(
        "`[修改文档] ${selectedFile.name}",
        "`📎 ${selectedFile.name}"
    )

    # 2. 替换 formData 参数名
    html = html.replace(
        "formData.append('instruction', message);",
        "formData.append('message', message);"
    )
    html = html.replace(
        "formData.append('username', currentUser);",
        "formData.append('session_id', currentChatId);"
    )

    # 3. 替换请求地址
    html = html.replace(
        "/api/v1/modify-document",
        "/api/v1/chat-with-file"
    )

    # 4. 替换返回处理逻辑
    html = html.replace(
        """let content = data.message || '修改完成';
            if (data.download_url) content += `\\n\\n<a class="download-link" href="${data.download_url}" download>⬇️ 下载修改后的文件</a>`;
            addMessageToUI('assistant', data.success ? content : '修改失败: ' + data.message);""",
        """if (data.response) {
                addMessageToUI('assistant', data.response);
            } else {
                addMessageToUI('assistant', data.detail || '处理失败');
            }
            await loadChatList();"""
    )

    # 5. 更新 onFileSelected - 添加动态 placeholder
    html = html.replace(
        "document.getElementById('fileBar').style.display = 'flex'; }",
        "document.getElementById('fileBar').style.display = 'flex'; document.getElementById('msgInput').placeholder = '针对此文件输入问题，或修改要求...'; }"
    )

    # 6. 更新 removeFile - 恢复 placeholder
    html = html.replace(
        "document.getElementById('fileBar').style.display = 'none'; }",
        "document.getElementById('fileBar').style.display = 'none'; document.getElementById('msgInput').placeholder = '输入问题...'; }"
    )

    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)
    print("  OK app/static/index.html")

# ========== 8. requirements.txt ==========
req_path = os.path.join(BASE, "requirements.txt")
with open(req_path, "r", encoding="utf-8") as f:
    req = f.read()
if "fpdf2" not in req:
    # 在 docx2txt 行后面添加 fpdf2
    req = req.replace("docx2txt>=0.8", "docx2txt>=0.8\nfpdf2>=2.7.0")
    with open(req_path, "w", encoding="utf-8") as f:
        f.write(req)
    print("  OK requirements.txt 添加 fpdf2")
else:
    print("  OK requirements.txt 已包含 fpdf2")

print("\n" + "=" * 50)
print("补丁 v2 完成！")
print()
print("下一步操作：")
print("1. 安装 fpdf2:")
print("   pip install fpdf2 -i https://pypi.tuna.tsinghua.edu.cn/simple")
print()
print("2. 【重要】修复 ChromaDB DLL 报错（二选一）：")
print("   方案A - 降级 chromadb（推荐，最稳）：")
print("   pip install chromadb==0.4.24 --force-reinstall -i https://pypi.tuna.tsinghua.edu.cn/simple")
print("   方案B - 不降级，当前代码已兼容 ChromaDB 不可用的情况")
print()
print("3. 重启服务：")
print("   python -m uvicorn app.main:app --host 0.0.0.0 --port 8000")
print()
print("4. 浏览器按 Ctrl+F5 强制刷新（清除缓存）")
print("=" * 50)
